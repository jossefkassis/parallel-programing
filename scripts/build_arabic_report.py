from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "report-ar.md"
OUTPUT = ROOT / "docs" / "report-ar.docx"
ARABIC_DIGITS = str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩")


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            r_pr.append(rtl)


def style_run(run, size=12, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", style=None, size=12, bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold)
    set_rtl(p)
    return p


def localize_prefix_number(text):
    if ". " not in text:
        return text
    prefix, rest = text.split(". ", 1)
    if prefix.isdigit():
        return prefix.translate(ARABIC_DIGITS) + ". " + rest
    return text


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            doc.add_paragraph()
            continue
        if line == "---":
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            add_para(doc, line[2:], "Heading 1", size=18, bold=True)
        elif line.startswith("## "):
            add_para(doc, localize_prefix_number(line[3:]), "Heading 2", size=15, bold=True)
        elif line.startswith("### "):
            add_para(doc, line[4:], "Heading 3", size=13, bold=True)
        elif in_code:
            p = add_para(doc, line, size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("- "):
            add_para(doc, "• " + line[2:], size=12)
        elif len(line) > 2 and line[0].isdigit() and line[1:3] == ". ":
            add_para(doc, localize_prefix_number(line), size=12)
        else:
            add_para(doc, line, size=12)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
